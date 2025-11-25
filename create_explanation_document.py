#!/usr/bin/env python3
"""
Generate Word document explaining Ramadan & Eid calculation logic
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('How We Calculate Ramadan & Eid Effects', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Simple Explanation for Everyone')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(107, 114, 128)

doc.add_paragraph()

# Introduction
doc.add_heading('🎯 The Big Problem', 1)
p = doc.add_paragraph(
    "Imagine you're running a restaurant, and you want to know how much money you'll make next year. "
    "But there's a problem: Ramadan moves around the calendar every year!"
)

doc.add_heading('In 2025 (This Year):', 2)
doc.add_paragraph('Ramadan is in March (March 1-30)', style='List Bullet')
doc.add_paragraph('Eid holidays are at the end of March and beginning of April (March 31, April 1-3)', style='List Bullet')
doc.add_paragraph('February is a normal month', style='List Bullet')

doc.add_heading('In 2026 (Next Year):', 2)
doc.add_paragraph('Ramadan starts earlier - in February (Feb 18-28) and continues into March (March 1-19)', style='List Bullet')
doc.add_paragraph('Eid holidays are in March (March 20-23)', style='List Bullet')
doc.add_paragraph('So February has both normal days AND Ramadan days!', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('This makes planning really hard! ')
run.bold = True
run.font.size = Pt(12)
run = p.add_run('How do we predict 2026 sales when the calendar is so different?')
run.font.size = Pt(12)

doc.add_page_break()

# Solution
doc.add_heading('🧮 Our Smart Solution', 1)
p = doc.add_paragraph('We use a "Copy and Paste" method, but in a smart way:')
p.runs[0].bold = True

doc.add_heading('Step 1: Sort Days into Buckets 🪣', 2)
doc.add_paragraph('First, we put every day into one of these buckets:')
doc.add_paragraph('🍽️ Normal Days - Regular business days', style='List Bullet')
doc.add_paragraph('🌙 Ramadan Days - People eat less during the day, big dinners at night', style='List Bullet')
doc.add_paragraph('🎉 Eid Days - Holiday celebration days with special patterns', style='List Bullet')

doc.add_heading('Step 2: Look at What Happened in 2025 👀', 2)
doc.add_paragraph('We look at our 2025 sales and calculate:')
doc.add_paragraph('Monday-Sunday averages for each type of day', style='List Bullet')
doc.add_paragraph('Like: "What\'s the average Monday during Ramadan?"', style='List Bullet')
doc.add_paragraph('Or: "What\'s the average Friday on a normal day?"', style='List Bullet')

doc.add_page_break()

# Example
doc.add_heading('📅 Example: Let\'s Plan February 2026', 1)

doc.add_heading('Part 1: Feb 1-17, 2026 (Normal Days)', 2)
p = doc.add_paragraph()
run = p.add_run('Question: ')
run.bold = True
p.add_run('What sales should we expect?')

p = doc.add_paragraph()
run = p.add_run('Answer: ')
run.bold = True
run.font.color.rgb = RGBColor(34, 139, 34)
p.add_run('Look at February 2025 (which was all normal days)')

doc.add_paragraph('Find all the Mondays in Feb 2025, average them → Use this for Mondays in Feb 1-17, 2026', style='List Bullet')
doc.add_paragraph('Find all the Tuesdays in Feb 2025, average them → Use this for Tuesdays in Feb 1-17, 2026', style='List Bullet')
doc.add_paragraph('...and so on for all 7 days of the week', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Why this works: ')
run.bold = True
run.font.color.rgb = RGBColor(59, 130, 246)
p.add_run('February weather, customer behavior, and business patterns are similar year to year')

doc.add_heading('Part 2: Feb 18-28, 2026 (Ramadan Days)', 2)
p = doc.add_paragraph()
run = p.add_run('Question: ')
run.bold = True
p.add_run('We never had Ramadan in February 2025! Where do we get the pattern?')

p = doc.add_paragraph()
run = p.add_run('Answer: ')
run.bold = True
run.font.color.rgb = RGBColor(34, 139, 34)
p.add_run('Look at March 2025 (which was all Ramadan days)')

doc.add_paragraph('Find all the Mondays during March Ramadan 2025, average them → Use for Mondays in Feb 18-28, 2026', style='List Bullet')
doc.add_paragraph('Find all the Tuesdays during March Ramadan 2025, average them → Use for Tuesdays in Feb 18-28, 2026', style='List Bullet')
doc.add_paragraph('...and so on', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Why this works: ')
run.bold = True
run.font.color.rgb = RGBColor(59, 130, 246)
p.add_run('Ramadan behavior is the same whether it\'s in February or March. People fast the same way!')

doc.add_page_break()

# Eid Special Case
doc.add_heading('🎉 Special Case: Eid Days Are Different!', 1)
p = doc.add_paragraph('Eid days are special - they\'re holidays that behave the same every year, no matter what day of the week they fall on.')
p.runs[0].bold = True

doc.add_heading('The Eid Copy Rule 📋', 2)
doc.add_paragraph('For 2026 Eid days, we just copy 2025 Eid sales exactly:')

# Create table for Eid mapping
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = '2025 Eid Day'
header_cells[1].text = '→ Copy To →'
header_cells[2].text = '2026 Eid Day'

# Data rows
mappings = [
    ('March 31, 2025 (1st Eid)', '→', 'March 20, 2026 (1st Eid)'),
    ('April 1, 2025 (2nd Eid)', '→', 'March 21, 2026 (2nd Eid)'),
    ('April 2, 2025 (3rd Eid)', '→', 'March 22, 2026 (3rd Eid)'),
    ('April 3, 2025 (4th Eid)', '→', 'March 23, 2026 (4th Eid)')
]

for i, mapping in enumerate(mappings, start=1):
    cells = table.rows[i].cells
    cells[0].text = mapping[0]
    cells[1].text = mapping[1]
    cells[2].text = mapping[2]

doc.add_paragraph()

doc.add_heading('Example:', 2)
doc.add_paragraph('If you made 1,000 BHD on 1st Eid Day in 2025', style='List Bullet')
doc.add_paragraph('You\'ll probably make 1,000 BHD on 1st Eid Day in 2026', style='List Bullet')
doc.add_paragraph('Even if one is Monday and the other is Friday!', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Why? ')
run.bold = True
run.font.color.rgb = RGBColor(59, 130, 246)
p.add_run('Because it\'s a holiday - people celebrate the same way regardless of which day of the week it is.')

doc.add_page_break()

# Complete Planning Calendar
doc.add_heading('🗓️ Complete 2026 Planning Calendar', 1)

doc.add_heading('February 2026:', 2)
p = doc.add_paragraph()
run = p.add_run('Days 1-17: ')
run.bold = True
p.add_run('Normal Days → Copy from Feb 2025 (normal days) using weekday averages')

p = doc.add_paragraph()
run = p.add_run('Days 18-28: ')
run.bold = True
p.add_run('Ramadan Days → Copy from March 2025 (Ramadan days) using weekday averages')

doc.add_heading('March 2026:', 2)
p = doc.add_paragraph()
run = p.add_run('Days 1-19: ')
run.bold = True
p.add_run('Ramadan Days → Copy from March 2025 (Ramadan days) using weekday averages')

p = doc.add_paragraph()
run = p.add_run('Days 20-23: ')
run.bold = True
p.add_run('Eid Days 🎉 → Copy EXACT values from March 31 + April 1-3, 2025')
doc.add_paragraph('1st Eid → 1st Eid', style='List Bullet 2')
doc.add_paragraph('2nd Eid → 2nd Eid', style='List Bullet 2')
doc.add_paragraph('3rd Eid → 3rd Eid', style='List Bullet 2')
doc.add_paragraph('4th Eid → 4th Eid', style='List Bullet 2')

p = doc.add_paragraph()
run = p.add_run('Days 24-31: ')
run.bold = True
p.add_run('Normal Days → Copy from Feb 2025 (normal days) using weekday averages')

doc.add_heading('April 2026:', 2)
p = doc.add_paragraph()
run = p.add_run('Days 1-30: ')
run.bold = True
p.add_run('Normal Days → Copy from April 4-30, 2025 (we skip April 1-3 because those were Eid in 2025) using weekday averages')

doc.add_page_break()

# Visual Example
doc.add_heading('🎨 Visual Example - March 20, 2026', 1)
doc.add_paragraph('Let\'s predict March 20, 2026 (Friday, 1st Eid Day):')

p = doc.add_paragraph()
run = p.add_run('❌ WRONG Way: ')
run.bold = True
run.font.color.rgb = RGBColor(220, 38, 38)
p.add_run('"Let\'s use Friday averages from March 2025"')
doc.add_paragraph('Why wrong? March 2025 Fridays were Ramadan days, not Eid', style='List Bullet 2')

p = doc.add_paragraph()
run = p.add_run('❌ WRONG Way: ')
run.bold = True
run.font.color.rgb = RGBColor(220, 38, 38)
p.add_run('"Let\'s use the same date - March 20, 2025"')
doc.add_paragraph('Why wrong? March 20, 2025 was a Ramadan day, not Eid', style='List Bullet 2')

p = doc.add_paragraph()
run = p.add_run('✅ CORRECT Way: ')
run.bold = True
run.font.color.rgb = RGBColor(34, 139, 34)
p.add_run('"Let\'s use March 31, 2025 (1st Eid Day)"')
doc.add_paragraph('March 31, 2025 was the 1st Eid Day', style='List Bullet 2')
doc.add_paragraph('It made 1,250 BHD (example)', style='List Bullet 2')
doc.add_paragraph('So March 20, 2026 should also make 1,250 BHD', style='List Bullet 2')
doc.add_paragraph('Even though March 31 was Monday and March 20 is Friday!', style='List Bullet 2')

p = doc.add_paragraph()
run = p.add_run('Why correct? ')
run.bold = True
run.font.color.rgb = RGBColor(59, 130, 246)
p.add_run('Because Eid behavior is the same every year - families celebrate, restaurants are busy, regardless of the weekday.')

doc.add_page_break()

# Summary
doc.add_heading('🧠 The Big Idea (Summary)', 1)
doc.add_paragraph('Think of it like this:')

doc.add_heading('1. Different days have different patterns:', 2)
doc.add_paragraph('Normal days = Regular school days', style='List Bullet')
doc.add_paragraph('Ramadan days = Special fasting days', style='List Bullet')
doc.add_paragraph('Eid days = Party/holiday days', style='List Bullet')

doc.add_heading('2. Weekdays matter for normal and Ramadan days:', 2)
doc.add_paragraph('Monday behavior ≠ Saturday behavior', style='List Bullet')
doc.add_paragraph('So we copy Monday → Monday, Saturday → Saturday', style='List Bullet')

doc.add_heading('3. Weekdays DON\'T matter for Eid:', 2)
doc.add_paragraph('1st Eid is always special', style='List Bullet')
doc.add_paragraph('Doesn\'t matter if it\'s Monday or Saturday', style='List Bullet')
doc.add_paragraph('So we copy 1st Eid → 1st Eid, 2nd Eid → 2nd Eid', style='List Bullet')

doc.add_heading('4. We look for the same TYPE of day from 2025:', 2)
doc.add_paragraph('Need a 2026 normal day? → Find 2025 normal days', style='List Bullet')
doc.add_paragraph('Need a 2026 Ramadan day? → Find 2025 Ramadan days', style='List Bullet')
doc.add_paragraph('Need a 2026 Eid day? → Find 2025 Eid days', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('It\'s like having a recipe book:')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(124, 58, 237)

doc.add_paragraph('Want to make "Normal Friday"? → Look up "Normal Friday" recipe from 2025', style='List Bullet')
doc.add_paragraph('Want to make "Ramadan Monday"? → Look up "Ramadan Monday" recipe from 2025', style='List Bullet')
doc.add_paragraph('Want to make "1st Eid Day"? → Look up "1st Eid Day" recipe from 2025 (exactly!)', style='List Bullet')

doc.add_page_break()

# Why it works
doc.add_heading('📊 Why This System Works', 1)

doc.add_paragraph('Respects weekday patterns - Mondays behave like Mondays', style='List Number')
doc.add_paragraph('Respects Ramadan effects - Ramadan behavior stays consistent', style='List Number')
doc.add_paragraph('Respects Eid uniqueness - Each Eid day has its own character', style='List Number')
doc.add_paragraph('Uses real data - Everything comes from actual 2025 sales', style='List Number')
doc.add_paragraph('Adapts to calendar shifts - Works even when Ramadan moves months', style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Result: ')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(34, 139, 34)
run = p.add_run('Accurate 2026 predictions that account for the Islamic calendar shifting! 🎯')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(34, 139, 34)

doc.add_page_break()

# Important Note about TOTAL Calculation
doc.add_heading('📊 Important: How We Calculate TOTALS', 1)

p = doc.add_paragraph()
run = p.add_run('The 2026 TOTAL row shows the sum of actual daily values in the table, NOT a formula from the backend.')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(220, 38, 38)

doc.add_paragraph()

doc.add_heading('Why This Matters:', 2)
p = doc.add_paragraph(
    'When you look at the daily breakdown tables, the TOTAL at the bottom is calculated by adding up '
    'all the individual daily sales shown above it. This ensures what you see is what you get!'
)

doc.add_heading('Example: March 2026 TOTAL Calculation', 2)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('March 2026 Daily Values:')
run.bold = True
doc.add_paragraph('Days 1-19: Ramadan days (using weekday averages from March 2025)', style='List Bullet')
doc.add_paragraph('Days 20-23: Eid days (using EXACT values from March 31, April 1-3, 2025)', style='List Bullet')
doc.add_paragraph('Days 24-31: Normal days (using weekday averages from Feb 2025)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('TOTAL = ')
run.bold = True
p.add_run('Sum of all these daily values (Days 1 + 2 + 3 + ... + 31)')

doc.add_paragraph()

# Create calculation example table
calc_table = doc.add_table(rows=6, cols=2)
calc_table.style = 'Light List Accent 1'

calc_cells = calc_table.rows[0].cells
calc_cells[0].text = 'Component'
calc_cells[1].text = 'Value'

components = [
    ('Ramadan Days Total (Days 1-19)', '15,200 BHD'),
    ('Eid Days Total (Days 20-23)', '5,200 BHD (exact 2025 values)'),
    ('Normal Days Total (Days 24-31)', '6,400 BHD'),
    ('', ''),
    ('MARCH 2026 TOTAL', '26,800 BHD')
]

for i, (comp, val) in enumerate(components, start=1):
    cells = calc_table.rows[i].cells
    cells[0].text = comp
    cells[1].text = val
    if i == 5:  # Make TOTAL row bold
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_heading('This Approach Ensures:', 2)
doc.add_paragraph('✅ The TOTAL matches the sum of what you see in the table', style='List Bullet')
doc.add_paragraph('✅ Eid day values are included exactly as they appeared in 2025', style='List Bullet')
doc.add_paragraph('✅ Weekday averaging effects are properly reflected', style='List Bullet')
doc.add_paragraph('✅ No discrepancies between individual days and the monthly total', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Consistency Across All Formats:')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(59, 130, 246)

doc.add_paragraph('The same calculation method is used in:', style='List Bullet')
doc.add_paragraph('📱 Web interface (daily breakdown modal)', style='List Bullet 2')
doc.add_paragraph('📊 Excel exports', style='List Bullet 2')
doc.add_paragraph('📄 PDF reports', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('This guarantees that no matter where you look - web, Excel, or PDF - the numbers will always match! ')
run.font.color.rgb = RGBColor(34, 139, 34)
run.font.size = Pt(11)

# Save document
output_path = '/home/user/backend/Backend/Ramadan_Eid_Calculation_Explanation.docx'
doc.save(output_path)
print(f"✅ Word document created: {output_path}")
